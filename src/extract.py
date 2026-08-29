"""Extract text from the document archive into a jsonl corpus.

    python src/extract.py --parser pymupdf --out corpus2.jsonl --filter-mojibake
    python src/extract.py --parser pypdf   --out corpus.jsonl --no-filter-mojibake

Merges the former extract.py (pypdf, no mojibake filter) and extract2.py
(PyMuPDF, mojibake filter). Both PDF readers are kept; --parser picks one.

One line of jsonl per document: path, name, folder, text.
"""
import argparse
import json
import logging
import os
from pathlib import Path

import docx
import openpyxl
import pymupdf
from dotenv import load_dotenv
from pypdf import PdfReader

load_dotenv(Path(__file__).resolve().parents[1] / '.env')

RAG_DIR = Path(os.getenv('RAG_DIR', Path(__file__).resolve().parents[1]))

logging.getLogger("pypdf").setLevel(logging.ERROR)
logging.getLogger("pymupdf").setLevel(logging.ERROR)

WEIRD = 'ÀÁÂÃÄÅÆÇÈÉÊËÌÍÎÏÐÑÒÓÔÕÖØÙÚÛÜÝÞßàáâãäåæçèéêëìíîïðñòóôõöøùúûüýþ²³¶µ§'


def is_mojibake(t):
    s = t[:3000]
    arm = sum(1 for c in s if '԰' <= c <= '֏')
    cyr = sum(1 for c in s if 'Ѐ' <= c <= 'ӿ')
    weird = sum(1 for c in s if c in WEIRD)
    return weird > 40 and weird > (arm + cyr)


def from_pdf_pymupdf(p):
    parts = []
    d = pymupdf.open(p)
    for page in d:
        parts.append(page.get_text())
    d.close()
    return "\n".join(parts)


def from_pdf_pypdf(p):
    return "\n".join((pg.extract_text() or "") for pg in PdfReader(p).pages)


def from_docx(p):
    d = docx.Document(p)
    parts = [par.text for par in d.paragraphs if par.text.strip()]
    for t in d.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def from_xlsx(p):
    wb = openpyxl.load_workbook(p, data_only=True, read_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"### Лист: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


PARSERS = {'pymupdf': from_pdf_pymupdf, 'pypdf': from_pdf_pypdf}


def extract(root, out_path, from_pdf, filter_mojibake):
    handlers = {".pdf": from_pdf, ".docx": from_docx, ".xlsx": from_xlsx}
    ok = skip = fail = moji = 0
    with open(out_path, "w", encoding="utf-8") as out:
        for dirpath, _, files in os.walk(root):
            for f in files:
                ext = os.path.splitext(f)[1].lower()
                if ext not in handlers:
                    continue
                path = os.path.join(dirpath, f)
                try:
                    text = handlers[ext](path)
                    if len(text.strip()) < 100:
                        skip += 1
                        continue
                    if filter_mojibake and is_mojibake(text):
                        moji += 1
                        continue
                    out.write(json.dumps({
                        "path": path,
                        "name": f,
                        "folder": os.path.relpath(dirpath, root),
                        "text": text
                    }, ensure_ascii=False) + "\n")
                    ok += 1
                    if ok % 200 == 0:
                        print("обработано:", ok, flush=True)
                except Exception:
                    fail += 1
    return ok, skip, moji, fail


def main():
    ap = argparse.ArgumentParser(description='Extract the archive into a jsonl corpus.')
    ap.add_argument('--parser', choices=sorted(PARSERS), default='pymupdf',
                    help='PDF reader (default: pymupdf)')
    ap.add_argument('--out', required=True,
                    help='output jsonl, relative to RAG_DIR')
    ap.add_argument('--filter-mojibake', dest='filter_mojibake',
                    action=argparse.BooleanOptionalAction, default=True,
                    help='drop legacy-encoding documents (default: on)')
    args = ap.parse_args()

    root = os.getenv('ARCHIVE_ROOT')
    if not root:
        raise SystemExit('ARCHIVE_ROOT is not set. Copy .env.example to .env and fill it in.')

    out_path = Path(args.out)
    if not out_path.is_absolute():
        out_path = RAG_DIR / out_path

    print(f'archive  : {root}')
    print(f'out      : {out_path}')
    print(f'parser   : {args.parser}')
    print(f'mojibake : {"filtered" if args.filter_mojibake else "kept"}')
    print(flush=True)

    ok, skip, moji, fail = extract(root, out_path, PARSERS[args.parser],
                                   args.filter_mojibake)
    print(f"\nГотово. Извлечено: {ok}, пустых: {skip}, mojibake: {moji}, ошибок: {fail}", flush=True)


if __name__ == '__main__':
    main()
